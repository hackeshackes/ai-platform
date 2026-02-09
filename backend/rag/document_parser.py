"""
Multi-Format Document Parser - RAG Enhancement

支持PDF、Word、Markdown等多种格式的文档解析
"""
import os
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from enum import Enum
import hashlib
import json


class DocumentFormat(Enum):
    """支持的文档格式"""
    PDF = "pdf"
    WORD = "docx"
    MARKDOWN = "md"
    TEXT = "txt"
    HTML = "html"
    UNKNOWN = "unknown"


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    doc_id: str
    file_name: str
    format: DocumentFormat
    title: str
    content: str
    sections: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    created_at: str
    chunk_count: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "doc_id": self.doc_id,
            "file_name": self.file_name,
            "format": self.format.value,
            "title": self.title,
            "content": self.content,
            "sections": self.sections,
            "metadata": self.metadata,
            "created_at": self.created_at,
            "chunk_count": self.chunk_count
        }


class BaseParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文档"""
        pass
    
    @abstractmethod
    def get_supported_formats(self) -> List[DocumentFormat]:
        """获取支持的格式"""
        pass
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        file_hash = hashlib.md5(
            f"{file_path}{os.path.getmtime(file_path)}".encode()
        ).hexdigest()[:12]
        return f"doc_{Path(file_path).stem}_{file_hash}"
    
    def _chunk_text(
        self, 
        text: str, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200
    ) -> List[str]:
        """文本分块"""
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = min(start + chunk_size, text_len)
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - chunk_overlap
            
            if start >= text_len:
                break
                
        return chunks


class MarkdownParser(BaseParser):
    """Markdown文档解析器"""
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        return [DocumentFormat.MARKDOWN, DocumentFormat.TEXT]
    
    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 解析Markdown结构
        sections = self._parse_markdown_structure(content)
        
        # 提取标题
        title = self._extract_title(content)
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            file_name=os.path.basename(file_path),
            format=DocumentFormat.MARKDOWN,
            title=title,
            content=content,
            sections=sections,
            metadata={
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "encoding": "utf-8"
            },
            created_at=self._get_timestamp()
        )
    
    def _parse_markdown_structure(self, content: str) -> List[Dict[str, Any]]:
        """解析Markdown标题结构"""
        sections = []
        lines = content.split('\n')
        current_section = {"level": 0, "title": "Introduction", "content": []}
        
        for line in lines:
            if line.startswith('#'):
                if current_section["content"]:
                    sections.append(current_section)
                level = len(line.split()[0])
                title = line[line.find(' ')+1:].strip()
                current_section = {
                    "level": level,
                    "title": title,
                    "content": []
                }
            else:
                current_section["content"].append(line)
        
        if current_section["content"]:
            sections.append(current_section)
        
        return sections
    
    def _extract_title(self, content: str) -> str:
        """提取文档标题"""
        for line in content.split('\n'):
            if line.startswith('# '):
                return line[2:].strip()
        return Path(file_path).stem
    
    def _get_timestamp(self) -> str:
        from datetime import datetime
        return datetime.utcnow().isoformat()


class PDFParser(BaseParser):
    """PDF文档解析器（简化实现）"""
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        return [DocumentFormat.PDF]
    
    def parse(self, file_path: str) -> ParsedDocument:
        # 尝试使用PyPDF2或其他PDF库
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() + "\n"
                
                metadata = {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                    "page_count": len(reader.pages),
                    "author": reader.metadata.author if reader.metadata else None,
                    "title": reader.metadata.title if reader.metadata else None
                }
                
        except ImportError:
            # 如果没有PyPDF2，返回占位实现
            content = f"[PDF解析需要安装PyPDF2] {file_path}"
            metadata = {
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "page_count": 0,
                "note": "Install PyPDF2 for full PDF parsing"
            }
        
        title = metadata.get("title", Path(file_path).stem)
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            file_name=os.path.basename(file_path),
            format=DocumentFormat.PDF,
            title=title,
            content=content,
            sections=[{"level": 1, "title": title, "content": content.split('\n')}],
            metadata=metadata,
            created_at=self._get_timestamp()
        )
    
    def _get_timestamp(self) -> str:
        from datetime import datetime
        return datetime.utcnow().isoformat()


class WordParser(BaseParser):
    """Word文档解析器（简化实现）"""
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        return [DocumentFormat.WORD]
    
    def parse(self, file_path: str) -> ParsedDocument:
        # 尝试使用python-docx
        try:
            import docx
            
            doc = docx.Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
            
            metadata = {
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
        except ImportError:
            content = f"[Word解析需要安装python-docx] {file_path}"
            metadata = {
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "note": "Install python-docx for full Word parsing"
            }
        
        title = Path(file_path).stem
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            file_name=os.path.basename(file_path),
            format=DocumentFormat.WORD,
            title=title,
            content=content,
            sections=[{"level": 1, "title": title, "content": content.split('\n')}],
            metadata=metadata,
            created_at=self._get_timestamp()
        )
    
    def _get_timestamp(self) -> str:
        from datetime import datetime
        return datetime.utcnow().isoformat()


class TextParser(BaseParser):
    """纯文本解析器"""
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        return [DocumentFormat.TEXT]
    
    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            file_name=os.path.basename(file_path),
            format=DocumentFormat.TEXT,
            title=Path(file_path).stem,
            content=content,
            sections=[{"level": 1, "title": "Content", "content": content.split('\n')}],
            metadata={
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "encoding": "utf-8"
            },
            created_at=self._get_timestamp()
        )
    
    def _get_timestamp(self) -> str:
        from datetime import datetime
        return datetime.utcnow().isoformat()


class DocumentParser:
    """文档解析器工厂"""
    
    _parsers: Dict[DocumentFormat, BaseParser] = {
        DocumentFormat.MARKDOWN: MarkdownParser(),
        DocumentFormat.PDF: PDFParser(),
        DocumentFormat.WORD: WordParser(),
        DocumentFormat.TEXT: TextParser()
    }
    
    @classmethod
    def get_parser(cls, file_path: str) -> BaseParser:
        """根据文件路径获取合适的解析器"""
        ext = Path(file_path).suffix.lower()
        
        format_map = {
            '.pdf': DocumentFormat.PDF,
            '.docx': DocumentFormat.WORD,
            '.md': DocumentFormat.MARKDOWN,
            '.txt': DocumentFormat.TEXT,
            '.html': DocumentFormat.HTML
        }
        
        doc_format = format_map.get(ext, DocumentFormat.UNKNOWN)
        
        if doc_format in cls._parsers:
            return cls._parsers[doc_format]
        
        return cls._parsers[DocumentFormat.TEXT]
    
    @classmethod
    def parse(cls, file_path: str) -> ParsedDocument:
        """解析文档"""
        parser = cls.get_parser(file_path)
        return parser.parse(file_path)
    
    @classmethod
    def parse_batch(cls, file_paths: List[str]) -> List[ParsedDocument]:
        """批量解析文档"""
        return [cls.parse(fp) for fp in file_paths]
    
    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """获取支持的格式列表"""
        return [fmt.value for fmt in DocumentFormat if fmt != DocumentFormat.UNKNOWN]
